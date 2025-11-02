from PyPDF2 import PdfReader
from docx import Document # Requires python-docx
import os
import fitz  # PyMuPDF
import io
from PIL import Image
import pytesseract
import cv2
import numpy as np
import re

# --- IMPORTANT ---
# On Windows, you may need to set this path to where Tesseract is installed.
tesseract_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
if os.path.exists(tesseract_path):
    pytesseract.pytesseract.tesseract_cmd = tesseract_path

def preprocess_image_for_ocr(image_bytes):
    """Converts image to grayscale and applies thresholding to improve OCR accuracy."""
    try:
        # Convert bytes to numpy array
        nparr = np.frombuffer(image_bytes, np.uint8)
        # Decode image
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        
        # Convert to grayscale
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        # Apply a binary threshold to get a black and white image
        # This can help in separating text from the background
        _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY)
        
        # Convert back to bytes for Pillow
        is_success, buffer = cv2.imencode(".png", thresh)
        if is_success:
            return buffer.tobytes()
    except Exception as e:
        print(f"Image preprocessing error: {e}")
    # Return original bytes if preprocessing fails
    return image_bytes

def clean_ocr_text(text):
    """Removes common OCR noise and non-meaningful characters."""
    # Remove lines that are mostly non-alphanumeric characters
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        # Count alphanumeric characters
        alnum_count = sum(c.isalnum() for c in line)
        # If the line is short and has few alphanumeric chars, it's likely noise
        if len(line.strip()) > 2 and alnum_count / len(line.strip()) > 0.5:
            cleaned_lines.append(line)
    
    text = '\n'.join(cleaned_lines)

    # Remove specific noisy patterns like "Made with GAMMA"
    text = re.sub(r'Made with GAMMA', '', text, flags=re.IGNORECASE)
    
    # Remove any remaining lines that are just whitespace
    return '\n'.join(line for line in text.split('\n') if line.strip()).strip()


def perform_ocr(image_bytes):
    """Performs OCR on image bytes and returns the extracted text."""
    try:
        # Preprocess the image to improve OCR quality
        processed_image_bytes = preprocess_image_for_ocr(image_bytes)
        
        image = Image.open(io.BytesIO(processed_image_bytes))
        
        # The user can specify a language for OCR, e.g., lang='eng+fra'
        raw_text = pytesseract.image_to_string(image)
        
        # Clean the extracted text to remove noise
        cleaned_text = clean_ocr_text(raw_text)
        
        return cleaned_text
    except Exception as e:
        print(f"OCR Error: {e}")
        return ""

def get_upload_dir():
    """Helper to get the temporary uploads directory path."""
    # Assumes the temp_uploads folder is at the same level as the backend folder
    return os.path.join(os.getcwd(), '..', 'temp_uploads')

def extract_images_from_pdf(file_path, task_id):
    """Extracts images from a PDF, saves them, and performs OCR."""
    image_data = []
    try:
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            for img_index, img in enumerate(doc.get_page_images(page_num)):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                
                # Save the image
                image_filename = f"{task_id}_page{page_num+1}_img{img_index+1}.png"
                image_path = os.path.join(get_upload_dir(), image_filename)
                
                with open(image_path, "wb") as f:
                    f.write(image_bytes)

                # Perform OCR
                alt_text = perform_ocr(image_bytes)
                image_data.append({'filename': image_filename, 'alt_text': alt_text})

        print(f"Extracted and OCR'd {len(image_data)} images from PDF.")
    except Exception as e:
        print(f"PDF Image Extraction Error: {e}")
    return image_data

def extract_images_from_docx(file_path, task_id):
    """Extracts images from a DOCX, saves them, and performs OCR."""
    image_data = []
    try:
        doc = Document(file_path)
        for i, rel in enumerate(doc.part.rels.values()):
            if "image" in rel.target_ref:
                image_bytes = rel.target_part.blob
                
                # Save the image
                image_filename = f"{task_id}_img{i+1}.png"
                image_path = os.path.join(get_upload_dir(), image_filename)

                with open(image_path, "wb") as f:
                    f.write(image_bytes)

                # Perform OCR
                alt_text = perform_ocr(image_bytes)
                image_data.append({'filename': image_filename, 'alt_text': alt_text})

        print(f"Extracted and OCR'd {len(image_data)} images from DOCX.")
    except Exception as e:
        print(f"DOCX Image Extraction Error: {e}")
    return image_data

def extract_pdf_text(file_path):
    """Extracts raw text from a PDF file."""
    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            text += page.extract_text() or ""
    except Exception as e:
        print(f"PDF Parsing error: {e}")
        return None
    return text.strip()

def extract_docx_text(file_path):
    """Extracts raw text from a DOCX file."""
    text = []
    try:
        document = Document(file_path)
        for paragraph in document.paragraphs:
            # Preserve basic structure: treat each paragraph as a block
            text.append(paragraph.text)
    except Exception as e:
        print(f"DOCX Parsing error: {e}")
        return None
    # Join paragraphs with two newlines for readability/separation
    return '\n\n'.join(text).strip()

def extract_text_and_images(file_path, task_id):
    """
    Detects file type and calls appropriate extractors for both text and images.
    Returns text and a list of dictionaries containing image filenames and alt text.
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    text = ""
    image_data = []
    file_type = "unknown"

    if file_extension == '.pdf':
        file_type = "pdf"
        text = extract_pdf_text(file_path)
        image_data = extract_images_from_pdf(file_path, task_id)
    elif file_extension == '.docx':
        file_type = "docx"
        text = extract_docx_text(file_path)
        image_data = extract_images_from_docx(file_path, task_id)
    
    # Combine extracted text with OCR'd text from images
    full_text = [text]
    for item in image_data:
        if item['alt_text']:
            full_text.append(f"\n\n--- Image Description ---\n{item['alt_text']}\n--- End Image Description ---\n")

    return ''.join(full_text), image_data, file_type


def extract_text(file_path):
    """Detects file type and calls the appropriate extractor."""
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return extract_pdf_text(file_path), "pdf"
    elif file_extension == '.docx':
        return extract_docx_text(file_path), "docx"
    else:
        # Handle other types if necessary, though MVP focuses on these
        return "Unsupported file type.", "unknown"
